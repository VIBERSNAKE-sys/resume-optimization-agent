"""
DOCX Generator — Resume document assembly using python-docx.

Generates the final resume .docx file from assembled D2 content.
Handles formatting, styles, margins, and section structure.
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")


def set_margins(doc: "Document", top: float = 0.75, bottom: float = 0.75,
                left: float = 0.75, right: float = 0.75):
    """Set page margins in inches."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_section_header(doc: "Document", text: str):
    """Add a bolded section header with spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)

    # Add a thin bottom border to the paragraph
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), '000000')
    p_bdr.append(bottom_bdr)
    p_pr.append(p_bdr)

    return p


def add_name_header(doc: "Document", name: str, contact: str, links: str = ""):
    """Add candidate name and contact information."""
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(14)

    # Contact
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    contact_run = contact_para.add_run(contact)
    contact_run.font.size = Pt(10)

    # Links (if any)
    if links:
        links_para = doc.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.paragraph_format.space_after = Pt(4)
        links_run = links_para.add_run(links)
        links_run.font.size = Pt(10)


def add_summary(doc: "Document", summary_text: str):
    """Add the professional summary section."""
    add_section_header(doc, "Professional Summary")
    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)


def add_skills(doc: "Document", skills_text: str):
    """Add the skills section."""
    add_section_header(doc, "Skills")
    p = doc.add_paragraph(skills_text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)


def add_experience_role(doc: "Document", role: dict):
    """Add a single role to the experience section."""
    # Company + Location line
    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_before = Pt(6)
    company_para.paragraph_format.space_after = Pt(0)
    company_run = company_para.add_run(role.get("company", ""))
    company_run.bold = True
    company_run.font.size = Pt(11)

    location = role.get("location", "")
    if location:
        location_run = company_para.add_run(f" | {location}")
        location_run.font.size = Pt(11)

    # Title + Dates line
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run(role.get("title", ""))
    title_run.font.size = Pt(11)

    dates = role.get("dates", "")
    if dates:
        # Right-align the dates
        tab_stop = OxmlElement('w:tab')
        title_para._p.append(tab_stop)
        dates_run = title_para.add_run(f"\t{dates}")
        dates_run.font.size = Pt(11)

    # Bullets
    for bullet_text in role.get("bullets", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(bullet_text)
        run.font.size = Pt(11)


def add_project(doc: "Document", project: dict):
    """Add a single project entry."""
    # Title + Date
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(4)
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run(project.get("title", ""))
    title_run.bold = True
    title_run.font.size = Pt(11)

    date_range = project.get("date_range", "")
    if date_range:
        date_run = title_para.add_run(f" | {date_range}")
        date_run.font.size = Pt(11)

    # Objective, Action, Outcome (bold header format)
    for section_key, section_label in [("objective", "Objective"), ("action", "Action"), ("outcome", "Outcome")]:
        section_text = project.get(section_key, "")
        if section_text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            label_run = p.add_run(f"{section_label}: ")
            label_run.bold = True
            label_run.font.size = Pt(11)
            content_run = p.add_run(section_text)
            content_run.font.size = Pt(11)

    # Technology line
    tech = project.get("technology", "")
    if tech:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        tech_label = p.add_run("Technology: ")
        tech_label.italic = True
        tech_label.font.size = Pt(11)
        tech_content = p.add_run(tech)
        tech_content.font.size = Pt(11)


def add_education(doc: "Document", education_entries: list[dict]):
    """Add education section."""
    add_section_header(doc, "Education")
    for entry in education_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        degree_run = p.add_run(entry.get("degree", ""))
        degree_run.bold = True
        degree_run.font.size = Pt(11)

        institution = entry.get("institution", "")
        if institution:
            inst_run = p.add_run(f" | {institution}")
            inst_run.font.size = Pt(11)

        year = entry.get("year", "")
        if year:
            year_run = p.add_run(f" | {year}")
            year_run.font.size = Pt(11)


def generate_docx(resume_data: dict, output_path: str) -> str:
    """
    Generate a resume .docx file from structured resume data.

    resume_data structure:
    {
        "candidate": {"name": "", "contact": "", "links": ""},
        "summary": "text",
        "skills": "comma-separated text",
        "experience": [
            {
                "company": "", "location": "", "title": "", "dates": "",
                "bullets": ["bullet 1 text", "bullet 2 text"]
            }
        ],
        "projects": [
            {
                "title": "", "date_range": "",
                "objective": "", "action": "", "outcome": "", "technology": ""
            }
        ],
        "education": [{"degree": "", "institution": "", "year": ""}],
        "certifications": ["cert 1", "cert 2"],
        "volunteer": [
            {"organization": "", "title": "", "dates": "", "bullets": []}
        ]
    }
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()

    # Set margins
    set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75)

    # Set default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Header: Name + Contact
    candidate = resume_data.get("candidate", {})
    add_name_header(
        doc,
        name=candidate.get("name", ""),
        contact=candidate.get("contact", ""),
        links=candidate.get("links", ""),
    )

    # Professional Summary
    summary = resume_data.get("summary", "")
    if summary:
        add_summary(doc, summary)

    # Skills
    skills = resume_data.get("skills", "")
    if skills:
        add_skills(doc, skills)

    # Professional Experience
    experience = resume_data.get("experience", [])
    if experience:
        add_section_header(doc, "Professional Experience")
        for role in experience:
            add_experience_role(doc, role)

    # Independent Projects
    projects = resume_data.get("projects", [])
    if projects:
        add_section_header(doc, "Independent Projects")
        for project in projects:
            add_project(doc, project)

    # Education
    education = resume_data.get("education", [])
    if education:
        add_education(doc, education)

    # Certifications
    certifications = resume_data.get("certifications", [])
    if certifications:
        add_section_header(doc, "Certifications")
        for cert in certifications:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            cert_run = p.add_run(cert)
            cert_run.font.size = Pt(11)

    # Volunteer Experience
    volunteer = resume_data.get("volunteer", [])
    if volunteer:
        add_section_header(doc, "Volunteer Experience")
        for vol_role in volunteer:
            add_experience_role(doc, vol_role)

    # Save
    doc.save(output_path)
    return output_path


def generate_from_state_files(
    verification_report_path: str,
    human_decisions_path: str,
    draft_summary_path: str,
    draft_projects_path: str,
    draft_skills_path: str,
    jd_ingestion_path: str,
    output_path: str,
) -> str:
    """
    Generate resume from pipeline state files.
    This is the main entry point called by the Assembly Agent.
    """
    # Load state files
    with open(verification_report_path) as f:
        verification = json.load(f)
    with open(human_decisions_path) as f:
        decisions = json.load(f)
    with open(draft_summary_path) as f:
        summary_text = f.read().strip()
    with open(draft_projects_path) as f:
        projects_data = json.load(f)
    with open(draft_skills_path) as f:
        skills_data = json.load(f)
    with open(jd_ingestion_path) as f:
        ingestion = json.load(f)

    # Build experience section from verified bullets
    # Group bullets by role, in resume order
    role_order = []
    role_bullets = {}

    for bullet in verification.get("bullets", []):
        if bullet.get("verdict") == "PASS":
            role_id = bullet.get("role_id", "")
            if role_id not in role_bullets:
                role_bullets[role_id] = []
                role_order.append(role_id)
            role_bullets[role_id].append(bullet["bullet_text"])

    # Build experience list from ingestion report role order
    experience = []
    for role_info in ingestion.get("employment_continuity", {}).get("roles", []):
        role_id = role_info["role_id"]
        role_entry = {
            "company": role_info.get("company", role_id),
            "location": "",
            "title": role_info.get("title", ""),
            "dates": f"{role_info.get('start_date', '')} – {role_info.get('end_date', '')}",
            "bullets": role_bullets.get(role_id, ["[Bullets pending]"]),
        }
        experience.append(role_entry)

    # Build projects list
    projects = []
    for proj in projects_data.get("projects", []):
        if proj.get("decision") != "DROP":
            projects.append({
                "title": proj.get("title", ""),
                "date_range": proj.get("date_range", ""),
                "objective": proj.get("objective", ""),
                "action": proj.get("action", ""),
                "outcome": proj.get("outcome", ""),
                "technology": proj.get("technology", ""),
            })

    # Assemble resume data structure
    resume_data = {
        "candidate": {
            "name": "[CANDIDATE NAME]",  # To be filled from resume.txt
            "contact": "[Contact Info]",
            "links": "",
        },
        "summary": summary_text,
        "skills": skills_data.get("skills_text", ""),
        "experience": experience,
        "projects": projects,
        "education": [],
        "certifications": [],
        "volunteer": [],
    }

    return generate_docx(resume_data, output_path)


def main():
    """CLI usage: python docx_generator.py <resume_data.json> <output.docx>"""
    if len(sys.argv) < 3:
        print("Usage: python docx_generator.py <resume_data.json> <output.docx>")
        sys.exit(1)

    with open(sys.argv[1]) as f:
        resume_data = json.load(f)

    output_path = sys.argv[2]
    result = generate_docx(resume_data, output_path)
    print(f"Resume generated: {result}")


if __name__ == "__main__":
    main()
